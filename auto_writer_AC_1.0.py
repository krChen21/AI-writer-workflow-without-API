import tkinter as tk
from tkinter import scrolledtext, messagebox, ttk
import threading
from playwright.sync_api import sync_playwright
import time
import os
from docx import Document  # 需安装: pip install python-docx

# ------------------ 自动化函数 (保持原样) ------------------

def ask_deepseek(prompt, status_callback, output_callback, new_chat=False, port=9222):
    """向 DeepSeek 提问并返回回复内容"""
    with sync_playwright() as p:
        try:
            browser = p.chromium.connect_over_cdp(f"http://localhost:{port}")
            status_callback("DeepSeek: 已连接浏览器")
            context = browser.contexts[0]
            page = None
            for existing_page in context.pages:
                if "deepseek.com" in existing_page.url:
                    page = existing_page
                    break
            if not page:
                status_callback("错误: 未找到 DeepSeek 标签页")
                return None

            page.bring_to_front()
            
            if new_chat:
                try:
                    new_chat_btn = page.locator("div[tabindex='0']:has(span:has-text('开启新对话'))")
                    if not new_chat_btn.count():
                        new_chat_btn = page.locator("div, button").filter(has_text="开启新对话").first
                    if new_chat_btn.count():
                        new_chat_btn.click()
                        page.wait_for_timeout(1000)
                        status_callback("新对话已创建（按钮）")
                    else:
                        status_callback("未找到新建对话按钮，继续执行（可能已在新建状态）")
                except Exception as e2:
                    status_callback(f"按钮点击也失败: {str(e2)[:30]}，继续执行")

            input_selectors = [
                "textarea[placeholder*='提问']",
                "textarea[placeholder*='消息']",
                "textarea[placeholder*='输入']",
                "textarea[placeholder*='发消息']",
                "textarea.ds-textarea",
                "textarea",
            ]
            textarea = None
            for selector in input_selectors:
                loc = page.locator(selector).first
                if loc.count():
                    try:
                        loc.wait_for(state="visible", timeout=3000)
                        textarea = loc
                        break
                    except:
                        pass
            if not textarea:
                status_callback("未找到输入框")
                return

            status_callback("正在输入提示词...")
            textarea.click()
            textarea.fill(prompt)
            page.evaluate("""
                (el) => {
                    el.dispatchEvent(new Event('input', { bubbles: true }));
                    el.dispatchEvent(new Event('change', { bubbles: true }));
                    el.focus();
                }
            """, textarea.element_handle())
            page.wait_for_timeout(300)

            status_callback("正在发送...")
            page.evaluate("""
                (el) => {
                    const keydown = new KeyboardEvent('keydown', { key: 'Enter', code: 'Enter', bubbles: true });
                    const keypress = new KeyboardEvent('keypress', { key: 'Enter', code: 'Enter', bubbles: true });
                    const keyup = new KeyboardEvent('keyup', { key: 'Enter', code: 'Enter', bubbles: true });
                    el.dispatchEvent(keydown);
                    el.dispatchEvent(keypress);
                    el.dispatchEvent(keyup);
                    const beforeInput = new InputEvent('beforeinput', { inputType: 'insertLineBreak', bubbles: true });
                    el.dispatchEvent(beforeInput);
                }
            """, textarea.element_handle())
            page.wait_for_timeout(500)

            if textarea.input_value().strip():
                status_callback("回车未清空，尝试点击发送按钮...")
                send_selectors = [
                    "button[aria-label='发送']",
                    "button[aria-label='Send']",
                    "button.ds-icon-button:not(.ds-icon-button--disabled)",
                    "button[type='submit']"
                ]
                for sel in send_selectors:
                    btn = page.locator(sel).first
                    if btn.count():
                        try:
                            btn.click(timeout=2000)
                            break
                        except:
                            pass

            status_callback("等待 AI 回复...")
            assistant_reply = ""

            stop_btn = page.locator("button:has-text('停止'), button:has-text('Stop')")
            try:
                stop_btn.wait_for(state="visible", timeout=30000)
                status_callback("生成中...")
                stop_btn.wait_for(state="hidden", timeout=180000)
                status_callback("生成完成")
            except:
                status_callback("监测内容变化...")
                last_text = ""
                stable = 0
                for _ in range(120):
                    msg = page.locator(".ds-markdown, .ds-message, .answer-content, .assistant-message").last
                    if msg.count():
                        current = msg.inner_text().strip()
                        if len(current) > 20:
                            if current == last_text:
                                stable += 1
                                if stable >= 3:
                                    break
                            else:
                                stable = 0
                                last_text = current
                    page.wait_for_timeout(1000)

            status_callback("读取回复内容...")
            reply = page.locator(".ds-markdown").last
            if not reply.count():
                reply = page.locator(".ds-message, .assistant-message, .answer-content").last
            if reply.count():
                assistant_reply = reply.inner_text().strip()
                if prompt in assistant_reply and len(assistant_reply) > len(prompt) + 10:
                    parts = assistant_reply.split(prompt)
                    if len(parts) > 1:
                        assistant_reply = parts[-1].strip()
            else:
                assistant_reply = "未能提取到回复"

            output_callback(f"【DeepSeek 回复】\n{assistant_reply}")
            status_callback("完成")
            return assistant_reply
        except Exception as e:
            status_callback(f"DeepSeek 出错: {str(e)[:30]}")
            return None

def ask_qianwen(prompt, status_callback, output_callback, new_chat=False, port=9222):
    """向通义千问提问并返回回复内容"""
    with sync_playwright() as p:
        try:
            browser = p.chromium.connect_over_cdp(f"http://localhost:{port}")
            status_callback("千问: 已连接浏览器")
            context = browser.contexts[0]
            page = None
            for existing_page in context.pages:
                if "qianwen" in existing_page.url or "tongyi" in existing_page.url:
                    page = existing_page
                    break
            if not page:
                status_callback("错误: 未找到通义千问标签页")
                return None

            page.bring_to_front()

            if new_chat:
                try:
                    new_chat_btn = page.locator("button:has-text('新建对话')")
                    if new_chat_btn.count():
                        status_callback("新建对话中...")
                        new_chat_btn.first.click()
                        page.wait_for_timeout(1000)
                        status_callback("新对话已创建")
                    else:
                        status_callback("未找到新建对话按钮，可能已在新对话中")
                except Exception as e:
                    status_callback(f"新建对话失败: {str(e)[:50]}，继续执行")
            else:
                status_callback("千问: 保持当前对话...")

            input_selectors = [
                "textarea[placeholder*='输入']",
                "textarea[placeholder*='提问']",
                "textarea[placeholder*='消息']",
                "div[contenteditable='true']",
                "textarea",
            ]
            textarea = None
            for selector in input_selectors:
                loc = page.locator(selector).first
                if loc.count():
                    try:
                        loc.wait_for(state="visible", timeout=3000)
                        textarea = loc
                        break
                    except:
                        pass
            if not textarea:
                status_callback("未找到输入框")
                return

            status_callback("正在输入提示词...")
            textarea.click()
            textarea.fill(prompt)
            page.evaluate("""
                (el) => {
                    el.dispatchEvent(new Event('input', { bubbles: true }));
                    el.dispatchEvent(new Event('change', { bubbles: true }));
                    el.focus();
                }
            """, textarea.element_handle())
            page.wait_for_timeout(300)

            status_callback("正在发送...")
            page.evaluate("""
                (el) => {
                    const keydown = new KeyboardEvent('keydown', { key: 'Enter', code: 'Enter', bubbles: true });
                    const keypress = new KeyboardEvent('keypress', { key: 'Enter', code: 'Enter', bubbles: true });
                    const keyup = new KeyboardEvent('keyup', { key: 'Enter', code: 'Enter', bubbles: true });
                    el.dispatchEvent(keydown);
                    el.dispatchEvent(keypress);
                    el.dispatchEvent(keyup);
                    const beforeInput = new InputEvent('beforeinput', { inputType: 'insertLineBreak', bubbles: true });
                    el.dispatchEvent(beforeInput);
                }
            """, textarea.element_handle())
            page.wait_for_timeout(500)

            try:
                is_empty = page.evaluate("(el) => el.innerText?.trim() === '' || el.value?.trim() === ''", textarea.element_handle())
            except:
                is_empty = True
            if not is_empty:
                status_callback("回车未清空，尝试点击发送按钮...")
                send_selectors = [
                    "button[aria-label='发送']",
                    "button[aria-label='Send']",
                    "button:has-text('发送')",
                    "button[type='submit']",
                    ".send-btn",
                    ".chat-panel-send-btn"
                ]
                for sel in send_selectors:
                    btn = page.locator(sel).first
                    if btn.count():
                        try:
                            btn.click(timeout=2000)
                            break
                        except:
                            pass

            status_callback("等待千问回复...")
            assistant_reply = ""

            stop_btn = page.locator("button:has-text('停止')")
            try:
                if stop_btn.count():
                    stop_btn.first.wait_for(state="visible", timeout=30000)
                    status_callback("生成中...")
                    stop_btn.first.wait_for(state="hidden", timeout=180000)
                    status_callback("生成完成（停止按钮消失）")
                else:
                    raise Exception("无停止按钮")
            except Exception as e:
                status_callback("未检测到停止按钮，使用内容稳定检测...")
                last_len = 0
                stable_count = 0
                required_stable = 5
                for _ in range(180):
                    cur_elem = page.locator(".qk-markdown").last
                    if not cur_elem.count():
                        page.wait_for_timeout(3000)
                        continue
                    cur_text = cur_elem.inner_text().strip()
                    cur_len = len(cur_text)
                    if cur_len > 20:
                        if cur_len == last_len:
                            stable_count += 1
                            if stable_count >= required_stable:
                                status_callback("生成完成（内容稳定）")
                                break
                        else:
                            stable_count = 0
                            last_len = cur_len
                    else:
                        stable_count = 0
                        last_len = cur_len
                    page.wait_for_timeout(1000)
                else:
                    status_callback("等待超时，强制读取")

            status_callback("读取回复内容...")
            assistant_reply = page.locator(".qk-markdown").last.inner_text().strip()
            if not assistant_reply:
                assistant_reply = "未能提取到回复内容"
            output_callback(f"【通义千问 回复】\n{assistant_reply}")
            status_callback("完成")
            return assistant_reply
        except Exception as e:
            status_callback(f"千问 出错: {str(e)[:30]}")
            return None

# ------------------ GUI 界面 (左:配置区, 右:上方对话记录+下方系统日志) ------------------

class AutoAskApp:
    def __init__(self, root):
        self.root = root
        root.title("AI 循环对谈 & 自动总结工具")
        root.geometry("1100x700")

        # 状态变量
        self.running = False
        self.continue_running = False
        self.full_history = []
        self.continue_history = []

        # ---- 风格预设数据结构 ----
        # 分类选项
        self.categories = ["通用", "男频", "女频"]
        # 类型字典
        self.types_dict = {
            "通用": ["玄幻", "科幻", "言情", "悬疑", "恐怖", "历史", "武侠", "奇幻", "轻小说", "同人", "现实", "末世"],
            "男频": ["玄幻", "仙侠", "都市", "科幻", "历史", "武侠", "游戏", "竞技体育", "末世求生", "悬疑惊悚", "轻小说"],
            "女频": ["古代言情", "现代言情", "穿越/重生", "玄幻言情", "科幻言情", "悬疑爱情", "校园青春", "轻喜剧"]
        }
        # 文风列表
        self.styles_list = [
            "简洁利落", "华丽细腻", "幽默诙谐", "黑暗冷峻", "诗意抒情",
            "口语纪实", "古风典雅", "港台腔", "翻译腔", "热血燃系", "悬疑压抑"
        ]

        # 使用 grid 布局，分为左右两列
        left_frame = tk.Frame(root, width=320, relief=tk.SUNKEN, borderwidth=1)
        left_frame.grid(row=0, column=0, sticky="ns", padx=5, pady=5)
        left_frame.grid_propagate(False)
        left_frame.config(width=320)

        right_frame = tk.Frame(root, relief=tk.SUNKEN, borderwidth=1)
        right_frame.grid(row=0, column=1, sticky="nsew", padx=5, pady=5)

        root.grid_columnconfigure(0, weight=0)
        root.grid_columnconfigure(1, weight=1)
        root.grid_rowconfigure(0, weight=1)

        # 右侧内部上下分列
        top_frame = tk.Frame(right_frame)
        top_frame.grid(row=0, column=0, sticky="nsew", padx=2, pady=2)
        bottom_frame = tk.Frame(right_frame)
        bottom_frame.grid(row=1, column=0, sticky="nsew", padx=2, pady=2)

        right_frame.grid_rowconfigure(0, weight=7)
        right_frame.grid_rowconfigure(1, weight=3)
        right_frame.grid_columnconfigure(0, weight=1)

        # ---- 左侧可滚动区域 ----
        left_canvas = tk.Canvas(left_frame, highlightthickness=0)
        left_scrollbar = tk.Scrollbar(left_frame, orient=tk.VERTICAL, command=left_canvas.yview)
        left_canvas.configure(yscrollcommand=left_scrollbar.set)
        left_scrollbar.pack(side=tk.RIGHT, fill=tk.Y)
        left_canvas.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)

        inner_left = tk.Frame(left_canvas)
        left_canvas.create_window((0, 0), window=inner_left, anchor=tk.NW, width=left_canvas.winfo_reqwidth())
        inner_left.bind("<Configure>", lambda e: left_canvas.configure(scrollregion=left_canvas.bbox("all")))

        def _on_mousewheel(event):
            left_canvas.yview_scroll(int(-1*(event.delta/120)), "units")
        left_canvas.bind_all("<MouseWheel>", _on_mousewheel)

        # 1. 初始提问
        tk.Label(inner_left, text="1. 初始提问 (主循环):", font=("Arial", 9, "bold")).pack(pady=(5,0), anchor=tk.W)
        self.initial_prompt_text = tk.Text(inner_left, height=2, font=("Arial", 9))
        self.initial_prompt_text.pack(fill=tk.X, padx=5, pady=3)
        self.initial_prompt_text.insert(tk.END, "请开始小说写作：主角是一名30多岁的男工程师，失业在家，在车库里搭建了一个关于电子、计算机、化学等一系列有条件实现的实验室，有很多人来拜访，其中发生了很多故事，改变了他的人生格局。")

        # ---- 新增：小说风格预设模块 ----
        style_frame = tk.LabelFrame(inner_left, text="小说风格预设 (可选)", padx=5, pady=3, font=("Arial", 9))
        style_frame.pack(fill=tk.X, padx=5, pady=5)

        # 分类选择
        cat_frame = tk.Frame(style_frame)
        cat_frame.pack(fill=tk.X, pady=2)
        tk.Label(cat_frame, text="分类:", font=("Arial", 8)).pack(side=tk.LEFT)
        self.category_var = tk.StringVar(value="通用")
        self.category_combo = ttk.Combobox(cat_frame, textvariable=self.category_var, values=self.categories, state="readonly", width=10)
        self.category_combo.pack(side=tk.LEFT, padx=5)
        self.category_combo.bind("<<ComboboxSelected>>", self.on_category_change)

        # 类型选择
        type_frame = tk.Frame(style_frame)
        type_frame.pack(fill=tk.X, pady=2)
        tk.Label(type_frame, text="类型:", font=("Arial", 8)).pack(side=tk.LEFT)
        self.type_var = tk.StringVar()
        self.type_combo = ttk.Combobox(type_frame, textvariable=self.type_var, values=self.types_dict["通用"], state="readonly", width=12)
        self.type_combo.pack(side=tk.LEFT, padx=5)

        # 文风多选 (Listbox 支持多选，限制最多3个)
        style_label_frame = tk.Frame(style_frame)
        style_label_frame.pack(fill=tk.X, pady=2)
        tk.Label(style_label_frame, text="文风 (最多可选3种):", font=("Arial", 8)).pack(anchor=tk.W)
        
        listbox_frame = tk.Frame(style_frame)
        listbox_frame.pack(fill=tk.BOTH, expand=True, pady=2)
        self.styles_listbox = tk.Listbox(listbox_frame, selectmode=tk.MULTIPLE, height=5, font=("Arial", 8))
        self.styles_listbox.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        scrollbar_styles = tk.Scrollbar(listbox_frame, orient=tk.VERTICAL, command=self.styles_listbox.yview)
        scrollbar_styles.pack(side=tk.RIGHT, fill=tk.Y)
        self.styles_listbox.config(yscrollcommand=scrollbar_styles.set)
        for style in self.styles_list:
            self.styles_listbox.insert(tk.END, style)
        self.styles_listbox.bind("<<ListboxSelect>>", self.on_style_select)

        # 提示标签
        self.style_warning_label = tk.Label(style_frame, text="", fg="red", font=("Arial", 7))
        self.style_warning_label.pack(anchor=tk.W)

        # 2. 对话前缀配置
        prefix_frame = tk.LabelFrame(inner_left, text="2. 对话前缀配置", padx=5, pady=3, font=("Arial", 9))
        prefix_frame.pack(fill=tk.X, padx=5, pady=3)
        tk.Label(prefix_frame, text="给DeepSeek的前穗 (基于千问评价):", font=("Arial", 8)).pack(anchor=tk.W)
        self.ds_prefix_entry = tk.Entry(prefix_frame, font=("Arial", 8))
        self.ds_prefix_entry.pack(fill=tk.X, pady=2)
        self.ds_prefix_entry.insert(0, "通义千问评价如下，请根据此评价继续写下一章：")
        tk.Label(prefix_frame, text="给千问的前缀 (基于DeepSeek创作):", font=("Arial", 8)).pack(anchor=tk.W)
        self.qw_prefix_entry = tk.Entry(prefix_frame, font=("Arial", 8))
        self.qw_prefix_entry.pack(fill=tk.X, pady=2)
        self.qw_prefix_entry.insert(0, "DeepSeek 创作如下，请你对其一致性、文笔风格一致性等进行批判性评价：")

        # 3. 总结配置
        summary_frame = tk.LabelFrame(inner_left, text="3. 结束总结配置 (仅主循环)", padx=5, pady=3, font=("Arial", 9))
        summary_frame.pack(fill=tk.X, padx=5, pady=3)
        tk.Label(summary_frame, text="总结提示词:", font=("Arial", 8)).pack(anchor=tk.W)
        self.summary_prompt_text = tk.Text(summary_frame, height=2, font=("Arial", 8))
        self.summary_prompt_text.pack(fill=tk.X, pady=3)
        self.summary_prompt_text.insert(tk.END, "请你根据以上所有的对话记录，整理并润色成一篇逻辑通顺、立意深刻的完整文章。")

        # 4. 主循环控制区
        ctrl_frame = tk.LabelFrame(inner_left, text="4. 主循环控制", padx=5, pady=3, font=("Arial", 9))
        ctrl_frame.pack(fill=tk.X, padx=5, pady=3)
        row_frame = tk.Frame(ctrl_frame)
        row_frame.pack(fill=tk.X)
        tk.Label(row_frame, text="轮数:", font=("Arial", 8)).pack(side=tk.LEFT)
        self.loop_count_entry = tk.Entry(row_frame, width=4, font=("Arial", 8))
        self.loop_count_entry.pack(side=tk.LEFT, padx=3)
        self.loop_count_entry.insert(0, "3")
        self.btn_start = tk.Button(row_frame, text="开始主循环", command=self.start_loop_thread, bg="#4CAF50", fg="white", font=("Arial", 8, "bold"), width=10)
        self.btn_start.pack(side=tk.LEFT, padx=3)
        self.btn_stop = tk.Button(row_frame, text="停止", command=self.stop_task, bg="#f44336", fg="white", font=("Arial", 8), width=5)
        self.btn_stop.pack(side=tk.LEFT, padx=3)

        # 5. 继续创作模块
        continue_frame = tk.LabelFrame(inner_left, text="5. 继续创作 (DeepSeek+千问交替)", padx=5, pady=3, font=("Arial", 9))
        continue_frame.pack(fill=tk.X, padx=5, pady=3)
        
        tk.Label(continue_frame, text="起始提示词:", font=("Arial", 8)).pack(anchor=tk.W)
        self.continue_start_prompt = tk.Text(continue_frame, height=2, width=50, font=("Arial", 8))
        self.continue_start_prompt.pack(fill=tk.X, pady=2)
        self.continue_start_prompt.insert(tk.END, "请继续续写小说，从上次结束的地方开始，保持风格和连贯性。")
        
        row2 = tk.Frame(continue_frame)
        row2.pack(fill=tk.X, pady=2)
        tk.Label(row2, text="轮数:", font=("Arial", 8)).pack(side=tk.LEFT)
        self.continue_round_entry = tk.Entry(row2, width=3, font=("Arial", 8))
        self.continue_round_entry.pack(side=tk.LEFT, padx=3)
        self.continue_round_entry.insert(0, "3")
        tk.Label(row2, text="端口:", font=("Arial", 8)).pack(side=tk.LEFT, padx=(5,0))
        self.port_entry = tk.Entry(row2, width=5, font=("Arial", 8))
        self.port_entry.pack(side=tk.LEFT, padx=3)
        self.port_entry.insert(0, "9222")
        
        btn_frame = tk.Frame(continue_frame)
        btn_frame.pack(fill=tk.X, pady=3)
        self.btn_continue_start = tk.Button(btn_frame, text="开始继续", command=self.start_continue_thread, bg="#2196F3", fg="white", font=("Arial", 8), width=8)
        self.btn_continue_start.pack(side=tk.LEFT, padx=2)
        self.btn_continue_stop = tk.Button(btn_frame, text="停止", command=self.stop_continue_task, bg="#ff9800", fg="white", font=("Arial", 8), width=5)
        self.btn_continue_stop.pack(side=tk.LEFT, padx=2)
        self.btn_export_continue = tk.Button(btn_frame, text="导出记录", command=self.export_continue_history, bg="#9C27B0", fg="white", font=("Arial", 8), width=8)
        self.btn_export_continue.pack(side=tk.LEFT, padx=2)
        
        info_label = tk.Label(continue_frame, text="说明: 每轮DeepSeek创作→千问评价→下一轮", 
                 fg="gray", font=("Arial", 7), wraplength=300, justify=tk.LEFT)
        info_label.pack(anchor=tk.W, pady=2)

        # ---- 右侧上方：对话流水记录 ----
        tk.Label(top_frame, text="对话流水记录:", font=("Arial", 9, "bold")).pack(anchor=tk.W, padx=5, pady=2)
        self.output_text = scrolledtext.ScrolledText(top_frame, font=("Arial", 9))
        self.output_text.pack(fill=tk.BOTH, expand=True, padx=5, pady=3)

        # ---- 右侧下方：系统日志 ----
        tk.Label(bottom_frame, text="系统日志:", font=("Arial", 9, "bold")).pack(anchor=tk.W, padx=5, pady=2)
        self.log_text = scrolledtext.ScrolledText(bottom_frame, font=("Arial", 7), fg="blue")
        self.log_text.pack(fill=tk.BOTH, expand=True, padx=5, pady=3)

        # 底部状态栏
        bottom_status = tk.Frame(root)
        bottom_status.grid(row=1, column=0, columnspan=2, sticky="ew", padx=5, pady=2)
        self.status_label = tk.Label(bottom_status, text="就绪", relief=tk.SUNKEN, anchor=tk.W, font=("Arial", 8))
        self.status_label.pack(fill=tk.X)

        # 初始化类型默认值
        self.on_category_change()

    # ---------- 风格预设辅助方法 ----------
    def on_category_change(self, event=None):
        """分类改变时更新类型下拉框"""
        cat = self.category_var.get()
        self.type_combo['values'] = self.types_dict.get(cat, [])
        if self.types_dict.get(cat):
            self.type_var.set(self.types_dict[cat][0])
        else:
            self.type_var.set("")

    def on_style_select(self, event=None):
        """限制文风最多选择3项"""
        selected = self.styles_listbox.curselection()
        if len(selected) > 3:
            # 取消最新选中的项
            last = selected[-1]
            self.styles_listbox.selection_clear(last)
            self.style_warning_label.config(text="最多只能选择3种文风！", fg="red")
            self.root.after(2000, lambda: self.style_warning_label.config(text=""))
        else:
            self.style_warning_label.config(text="")

    def build_style_prompt(self):
        """根据用户选择生成风格提示词，若无选择则返回空字符串"""
        cat = self.category_var.get()
        genre = self.type_var.get()
        selected_indices = self.styles_listbox.curselection()
        selected_styles = [self.styles_listbox.get(i) for i in selected_indices]
        
        if not genre and not selected_styles:
            return ""  # 没有任何选择
        
        parts = []
        if genre:
            parts.append(f"- 类型：{genre}")
        if selected_styles:
            parts.append(f"- 文风：{'、'.join(selected_styles)}")
        
        if parts:
            return "【小说创作要求】\n" + "\n".join(parts) + "\n\n请严格按照以上要求进行创作。"
        return ""

    # ---------- 原有方法 ----------
    def update_status(self, msg):
        self.status_label.config(text=f"状态: {msg}")
        timestamp = time.strftime("%H:%M:%S")
        self.log_text.insert(tk.END, f"[{timestamp}] {msg}\n")
        self.log_text.see(tk.END)
        self.root.update_idletasks()

    def update_output(self, text, role="", target_history="main"):
        self.output_text.insert(tk.END, text + "\n\n" + "-"*50 + "\n\n")
        self.output_text.see(tk.END)
        if target_history == "main":
            self.full_history.append(f"{role}:\n{text}")
        elif target_history == "continue":
            self.continue_history.append(f"{role}:\n{text}")

    def stop_task(self):
        self.running = False
        self.update_status("主循环停止")

    def stop_continue_task(self):
        self.continue_running = False
        self.update_status("继续创作已停止")
        self.btn_continue_start.config(state=tk.NORMAL)

    def export_to_word(self, filename, title, content_list):
        try:
            doc = Document()
            doc.add_heading(title, 0)
            for item in content_list:
                doc.add_paragraph(item)
                doc.add_paragraph("-" * 30)
            save_path = os.path.join(os.getcwd(), filename)
            doc.save(save_path)
            return save_path
        except Exception as e:
            return str(e)

    def export_continue_history(self):
        if not self.continue_history:
            messagebox.showwarning("无记录", "没有继续创作的记录可供导出。")
            return
        timestamp = time.strftime("%Y%m%d_%H%M%S")
        path = self.export_to_word(f"继续创作记录_{timestamp}.docx", "AI 继续创作记录", self.continue_history)
        messagebox.showinfo("导出成功", f"已导出至：{path}")

    # ------------------ 主循环 ------------------
    def start_loop_thread(self):
        if self.running:
            self.update_status("主循环已在运行中")
            return
        if self.continue_running:
            messagebox.showwarning("提示", "请先停止继续创作任务")
            return
        self.running = True
        self.full_history = []
        self.output_text.delete("1.0", tk.END)
        self.btn_start.config(state=tk.DISABLED)
        threading.Thread(target=self._run_loop_logic, daemon=True).start()

    def _run_loop_logic(self):
        try:
            loop_limit = int(self.loop_count_entry.get())
            base_prompt = self.initial_prompt_text.get("1.0", tk.END).strip()
            style_prompt = self.build_style_prompt()
            if style_prompt:
                # 将风格预设附加到初始提问后面
                current_input = f"{base_prompt}\n\n{style_prompt}"
                self.update_status("已附加小说风格预设到初始提问")
            else:
                current_input = base_prompt
            
            for i in range(loop_limit):
                if not self.running: break
                
                self.update_status(f"第 {i+1} 轮: 呼叫 DeepSeek")
                ds_reply = ask_deepseek(current_input, self.update_status, 
                                        lambda t: self.update_output(t, "DeepSeek", target_history="main"), 
                                        new_chat=(i==0))
                if not ds_reply: break
                
                qw_input = f"{self.qw_prefix_entry.get()}\n{ds_reply}"
                self.update_status(f"第 {i+1} 轮: 呼叫 通义千问")
                qw_reply = ask_qianwen(qw_input, self.update_status, 
                                       lambda t: self.update_output(t, "通义千问", target_history="main"), 
                                       new_chat=(i==0))
                if not qw_reply: break
                
                current_input = f"{self.ds_prefix_entry.get()}\n{qw_reply}"

            if self.running:
                self.update_status("正在进行最终总结...")
                summary_prompt_base = self.summary_prompt_text.get("1.0", tk.END).strip()
                full_context = "\n\n".join(self.full_history)
                final_summary_prompt = f"{summary_prompt_base}\n\n以下是原始对话记录：\n{full_context}"
                summary_result = ask_deepseek(final_summary_prompt, self.update_status, lambda t: None, new_chat=False)
                
                if summary_result:
                    self.update_output(summary_result, "【最终总结文章】", target_history="main")
                    timestamp = time.strftime("%Y%m%d_%H%M%S")
                    path1 = self.export_to_word(f"对话流水_{timestamp}.docx", "对话流水记录", self.full_history[:-1])
                    path2 = self.export_to_word(f"最终总结_{timestamp}.docx", "AI 总结文章", [summary_result])
                    messagebox.showinfo("完成", f"任务已结束！\n\n流水存档：{path1}\n总结存档：{path2}")
                else:
                    messagebox.showwarning("总结失败", "对话已结束，但总结过程出错。")
            self.update_status("所有任务已完成")
        except Exception as e:
            self.update_status(f"发生错误: {e}")
        finally:
            self.running = False
            self.btn_start.config(state=tk.NORMAL)

    # ------------------ 继续创作 ------------------
    def start_continue_thread(self):
        if self.continue_running:
            self.update_status("继续创作已在运行中")
            return
        if self.running:
            messagebox.showwarning("提示", "请先停止主对话循环")
            return
        
        try:
            rounds = int(self.continue_round_entry.get())
            if rounds <= 0:
                messagebox.showerror("错误", "轮数必须大于0")
                return
            port = int(self.port_entry.get())
            start_prompt = self.continue_start_prompt.get("1.0", tk.END).strip()
            if not start_prompt:
                messagebox.showerror("错误", "请输入起始提示词")
                return
        except ValueError:
            messagebox.showerror("错误", "请输入有效的数字")
            return
        
        self.continue_running = True
        self.btn_continue_start.config(state=tk.DISABLED)
        self.continue_history = []
        threading.Thread(target=self._run_continue_logic, args=(rounds, start_prompt, port), daemon=True).start()

    def _run_continue_logic(self, rounds, start_prompt, port):
        try:
            current_input = start_prompt
            
            for i in range(1, rounds + 1):
                if not self.continue_running:
                    break
                
                self.update_status(f"继续创作第 {i}/{rounds} 轮: DeepSeek")
                ds_reply = ask_deepseek(
                    current_input, 
                    self.update_status,
                    lambda t: self.update_output(t, f"DeepSeek(轮{i})", target_history="continue"),
                    new_chat=False,
                    port=port
                )
                if not ds_reply:
                    self.update_status(f"第{i}轮 DeepSeek 失败，停止")
                    break
                
                self.update_status(f"继续创作第 {i}/{rounds} 轮: 千问")
                qw_input = f"{self.qw_prefix_entry.get()}\n{ds_reply}"
                qw_reply = ask_qianwen(
                    qw_input,
                    self.update_status,
                    lambda t: self.update_output(t, f"通义千问(轮{i})", target_history="continue"),
                    new_chat=False,
                    port=port
                )
                if not qw_reply:
                    self.update_status(f"第{i}轮 千问失败，停止")
                    break
                
                current_input = f"{self.ds_prefix_entry.get()}\n{qw_reply}"
                time.sleep(1)
            
            if self.continue_running:
                self.update_status(f"继续创作完成，共 {rounds} 轮")
                messagebox.showinfo("完成", f"交替创作已完成 {rounds} 轮\n可点击「导出记录」保存。")
            else:
                self.update_status("继续创作已手动停止")
        except Exception as e:
            self.update_status(f"继续创作错误: {e}")
            messagebox.showerror("错误", f"继续创作出错: {e}")
        finally:
            self.continue_running = False
            self.btn_continue_start.config(state=tk.NORMAL)

if __name__ == "__main__":
    root = tk.Tk()
    app = AutoAskApp(root)
    root.mainloop()