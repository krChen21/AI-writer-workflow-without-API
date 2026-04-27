# -*- coding: utf-8 -*-
"""
工具函数模块：封装通用功能，如Word导出、状态/日志更新等
"""
import time
import os
import tkinter as tk
from tkinter import messagebox
from docx import Document  # 需安装: pip install python-docx


def update_status(status_label, log_text, msg):
    """
    更新界面状态和日志
    :param status_label: 状态标签组件
    :param log_text: 日志文本框组件
    :param msg: 要显示的消息
    """
    status_label.config(text=f"状态: {msg}")
    timestamp = time.strftime("%H:%M:%S")
    log_text.insert(tk.END, f"[{timestamp}] {msg}\n")
    log_text.see(tk.END)
    # 强制刷新界面
    status_label.master.update_idletasks()


def update_output(output_text, text, role="", history_list=None, history_type="main"):
    """
    更新对话输出框，并记录到历史列表
    :param output_text: 输出文本框组件
    :param text: 要显示的文本内容
    :param role: 角色标识（如DeepSeek/通义千问）
    :param history_list: 历史记录列表（主循环/继续创作）
    :param history_type: 历史类型（main/continue）
    """
    output_text.insert(tk.END, text + "\n\n" + "-"*50 + "\n\n")
    output_text.see(tk.END)
    
    # 记录到历史列表
    if history_list is not None:
        history_list.append(f"{role}:\n{text}")


def export_to_word(filename, title, content_list, output_dir):
    """
    将内容导出为Word文档
    :param filename: 保存的文件名
    :param title: 文档标题
    :param content_list: 要导出的内容列表
    :return: 成功返回保存路径，失败返回错误信息
    """
    try:
        doc = Document()
        # 添加标题
        doc.add_heading(title, 0)
        # 逐个添加内容
        for item in content_list:
            doc.add_paragraph(item)
            doc.add_paragraph("-" * 30)  # 分隔线
        # 保存路径（当前目录）

        if output_dir:
            # 确保目录存在
            os.makedirs(output_dir, exist_ok=True)
            save_path = os.path.join(output_dir, filename)
        else:
            save_path = os.path.join(os.getcwd(), filename)

        doc.save(save_path)
        return save_path
    except Exception as e:
        return str(e)


def export_history(history_list, output_dir ,export_type="continue"):
    """
    导出历史记录为Word（封装弹窗提示）
    :param history_list: 历史记录列表
    :param export_type: 导出类型（continue/主循环）
    """
    if not history_list:
        messagebox.showwarning("无记录", f"没有{export_type}创作的记录可供导出。")
        return
    
    # 生成带时间戳的文件名
    timestamp = time.strftime("%Y%m%d_%H%M%S")
    filename = f"{export_type}创作记录_{timestamp}.docx"
    title = f"AI {export_type}创作记录"
    
    # 执行导出
    path = export_to_word(filename, title, history_list, output_dir)
    
    # 提示结果
    if os.path.exists(path):
        messagebox.showinfo("导出成功", f"已导出至：{path}")
    else:
        messagebox.showerror("导出失败", f"导出出错：{path}")